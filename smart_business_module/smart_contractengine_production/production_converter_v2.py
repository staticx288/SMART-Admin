#!/usr/bin/env python3
"""
Production Floor SOP Converter V2
User-driven section selection and contract generation

Workflow:
1. Parse SOP → Extract numbered sections
2. User selects sections and assigns contract types
3. Generate YAML contracts from selected sections
"""

import re
from typing import Dict, List, Any, Optional
from pathlib import Path
from datetime import date


class ProductionConverterV2:
    """
    Enhanced converter with user-driven section selection
    """
    
    # Supported domains
    DOMAINS = [
        'LP', 'MPI', 'UT', 'RT', 'VT', 'WELD', 
        'MACHINE', 'HEAT', 'CP', 'PAINT', 'POLISH', 
        'ASSEMBLE', 'INSPECT'
    ]
    
    # Contract types
    CONTRACT_TYPES = [
        'SmartSafety',
        'SmartMaintenance',
        'SmartInventory',
        'SmartCompliance',
        'SmartSP',
        'SmartQA',
        'SmartStandards'
    ]
    
    def __init__(self):
        """Initialize converter"""
        pass
    
    def parse_sop(self, file_content: bytes, filename: str, domain: str) -> Dict[str, Any]:
        """
        Parse SOP and extract all sections for user selection
        
        Returns:
            {
                'domain': str,
                'filename': str,
                'sections': [
                    {'number': 1, 'title': 'Overview', 'content': '...'},
                    {'number': 4, 'title': 'Safety Precautions', 'content': '...'},
                    ...
                ]
            }
        """
        if domain not in self.DOMAINS:
            raise ValueError(f"Invalid domain: {domain}")
        
        # Parse document
        file_ext = Path(filename).suffix.lower()
        
        if file_ext == '.docx':
            document_text = self._parse_docx(file_content)
        elif file_ext == '.pdf':
            document_text = self._parse_pdf(file_content)
        elif file_ext in ['.md', '.txt']:
            document_text = file_content.decode('utf-8')
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
        
        # Extract numbered sections
        sections = self._extract_numbered_sections(document_text)
        
        return {
            'domain': domain,
            'filename': filename,
            'sections': sections,
            'total_sections': len(sections)
        }
    
    def _parse_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            from docx import Document
            from io import BytesIO
            
            doc = Document(BytesIO(file_content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Also get text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            
            return '\n'.join(paragraphs)
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {str(e)}")
    
    def _parse_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            from PyPDF2 import PdfReader
            from io import BytesIO
            
            pdf = PdfReader(BytesIO(file_content))
            text_parts = []
            
            for page in pdf.pages:
                text_parts.append(page.extract_text())
            
            return '\n'.join(text_parts)
        except Exception as e:
            raise ValueError(f"Failed to parse PDF: {str(e)}")
    
    def _extract_numbered_sections(self, document_text: str) -> List[Dict[str, Any]]:
        """
        Extract sections by looking for numbered headings like:
        1. Overview
        4. Safety Precautions
        10. Final Inspection
        """
        sections = []
        lines = document_text.split('\n')
        
        # Pattern to match numbered sections: "1. Title" or "1  Title" or "1) Title"
        section_pattern = re.compile(r'^(\d+)[\.)\s]+(.+)$')
        
        current_section = None
        current_content = []
        
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue
            
            # Check if this is a section header
            match = section_pattern.match(stripped)
            
            if match:
                # Save previous section
                if current_section is not None:
                    current_section['content'] = '\n'.join(current_content).strip()
                    sections.append(current_section)
                
                # Start new section
                section_num = int(match.group(1))
                section_title = match.group(2).strip()
                
                current_section = {
                    'number': section_num,
                    'title': section_title,
                    'content': ''
                }
                current_content = []
            elif current_section is not None:
                # Add line to current section content
                current_content.append(line)
        
        # Don't forget the last section
        if current_section is not None:
            current_section['content'] = '\n'.join(current_content).strip()
            sections.append(current_section)
        
        # Sort by section number
        sections.sort(key=lambda x: x['number'])
        
        return sections
    
    def generate_contract(
        self,
        section_content: str,
        contract_type: str,
        contract_id: str,
        domain: str,
        title: str = None
    ) -> Dict[str, Any]:
        """
        Generate a single YAML contract from a section
        
        Args:
            section_content: The extracted section text
            contract_type: One of CONTRACT_TYPES
            contract_id: User-provided ID (e.g., "PEN-SAFE-001")
            domain: Production domain
            title: Optional custom title
            
        Returns:
            {
                'id': str,
                'type': str,
                'filename': str,
                'yaml_content': str,
                'domain': str
            }
        """
        if contract_type not in self.CONTRACT_TYPES:
            raise ValueError(f"Invalid contract type: {contract_type}")
        
        # Extract steps from content
        steps = self._extract_steps(section_content, contract_type)
        
        # Generate title if not provided
        if not title:
            title = f"{domain} {contract_type} Contract"
        
        # Build YAML structure
        yaml_content = f"""smart_contract_id: {contract_id}
contract_type: {contract_type}
title: {title}
description: Extracted from SOP
status: active
created_date: '{date.today()}'
checklist_steps:
"""
        
        # Add steps
        for i, step in enumerate(steps, 1):
            yaml_content += f"  - step: {i}\n"
            yaml_content += f"    description: {step['description']}\n"
            yaml_content += f"    input_type: {step['input_type']}\n"
            
            # Add optional fields
            if 'unit' in step:
                yaml_content += f"    unit: {step['unit']}\n"
            if 'reference_value' in step:
                yaml_content += f"    reference_value: {step['reference_value']}\n"
        
        return {
            'id': contract_id,
            'type': contract_type,
            'filename': f"{contract_id}.yaml",
            'yaml_content': yaml_content,
            'domain': domain
        }
    
    def _extract_steps(self, content: str, contract_type: str) -> List[Dict[str, Any]]:
        """
        Extract checklist steps from section content
        Uses simple logic to identify action items
        """
        steps = []
        lines = content.split('\n')
        
        for line in lines:
            stripped = line.strip()
            if not stripped or len(stripped) < 10:
                continue
            
            # Look for numbered items, bullets, or sentences
            if any(stripped.startswith(c) for c in ['•', '-', '*', '1.', '2.', '3.']):
                step = {'description': stripped.lstrip('•-*123456789. '), 'input_type': 'checkbox'}
                
                # Check if it contains measurements
                if self._contains_measurement(stripped):
                    step['input_type'] = 'number'
                    unit = self._extract_unit(stripped)
                    if unit:
                        step['unit'] = unit
                    ref_value = self._extract_reference_value(stripped)
                    if ref_value:
                        step['reference_value'] = ref_value
                
                steps.append(step)
            elif re.match(r'^[A-Z]', stripped) and stripped.endswith('.'):
                # Complete sentences that look like instructions
                step = {'description': stripped, 'input_type': 'checkbox'}
                steps.append(step)
        
        # If no steps found, create one generic step
        if not steps:
            steps.append({
                'description': content[:200] + '...' if len(content) > 200 else content,
                'input_type': 'checkbox'
            })
        
        return steps
    
    def _contains_measurement(self, text: str) -> bool:
        """Check if text contains measurement indicators"""
        measurement_patterns = [
            r'\d+\s*(psi|ftc|°F|°C|µW/cm²|minutes|hours|mm|cm|inches)',
            r'(temperature|pressure|time|intensity|distance|flow)',
            r'(measure|record|log|monitor)'
        ]
        return any(re.search(pattern, text, re.IGNORECASE) for pattern in measurement_patterns)
    
    def _extract_unit(self, text: str) -> Optional[str]:
        """Extract unit from measurement text"""
        unit_pattern = r'\b(psi|ftc|°F|°C|µW/cm²|minutes|hours|mm|cm|inches)\b'
        match = re.search(unit_pattern, text, re.IGNORECASE)
        return match.group(1) if match else None
    
    def _extract_reference_value(self, text: str) -> Optional[str]:
        """Extract reference value from text like '≥ 40 psi' or '10-60 minutes'"""
        # Pattern for ranges: "10-60 minutes"
        range_pattern = r'(\d+[-–]\d+)\s*(?:psi|ftc|°F|°C|µW/cm²|minutes|hours|mm|cm|inches)?'
        match = re.search(range_pattern, text)
        if match:
            return match.group(1)
        
        # Pattern for comparisons: "≥ 40 psi"
        comparison_pattern = r'([≥≤<>]\s*\d+)\s*(?:psi|ftc|°F|°C|µW/cm²|minutes|hours|mm|cm|inches)?'
        match = re.search(comparison_pattern, text)
        if match:
            return match.group(1)
        
        return None


# For standalone testing
if __name__ == "__main__":
    print("🧪 Testing Production Converter V2...")
    
    converter = ProductionConverterV2()
    
    # Test with sample markdown
    sample_sop = """
# Penetrant Process Planning Form

1. Overview
This is the overview section.

4. Safety Precautions
- Wear nitrile gloves, safety glasses, and UV face shield
- Verify emergency shower is accessible
- Check ventilation is operational

8. Equipment
- Confirm UV-A black light intensity ≥ 1000 µW/cm²
- Ensure oven temperature 139–159°F
- Verify pressure gauge reads ≥ 40 psi

9. Process Steps
1. Clean parts with acetone
2. Apply penetrant via immersion (10-30 minutes dwell time)
3. Rinse with DI water at 40 psi max pressure
"""
    
    result = converter.parse_sop(sample_sop.encode(), 'test.md', 'LP')
    print(f"✅ Extracted {result['total_sections']} sections")
    
    for section in result['sections']:
        print(f"\nSection {section['number']}: {section['title']}")
    
    # Generate a sample contract
    if result['sections']:
        safety_section = next((s for s in result['sections'] if s['number'] == 4), None)
        if safety_section:
            contract = converter.generate_contract(
                section_content=safety_section['content'],
                contract_type='SmartSafety',
                contract_id='PEN-SAFE-001',
                domain='LP',
                title='Penetrant Booth Safety Pre-Check'
            )
            print(f"\n📄 Generated Contract:\n{contract['yaml_content']}")
    
    print("\n✅ Converter V2 test complete")
